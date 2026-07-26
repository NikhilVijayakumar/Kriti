"""generate_reference_docx.py — one-shot build script for reference/
template/Template_PCEMS2026_reference.docx, the pandoc --reference-doc
used by render-docx.py.

pandoc's default DOCX output uses Calibri; the official template
(reference/template/extracted/Template_PCEMS2026.txt) specifies Arial at
fixed sizes per heading level. python-docx can't set arbitrary style
templates pandoc will honor except through its documented --reference-doc
mechanism: pandoc maps its own Title/Heading 1/Heading 2/Heading 3/Body
Text (Normal) styles onto whatever a reference .docx defines for those
same style names. Rebuilding this one style set here (not hand-authoring
a .docx in Word) keeps it in version control and reproducible.

Run once from repo root: python script/render/generate_reference_docx.py
Not part of the runtime render pipeline — a build-time asset generator,
regenerate only if the template's formatting spec changes.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = (Path(__file__).resolve().parent.parent.parent / "reference" / "template"
            / "Template_PCEMS2026_reference.docx")


def _set_style_font(style, name="Arial", size=11, bold=False, italic=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic


def build():
    doc = Document()

    # Body text (pandoc's default HTML <p> maps to Normal)
    _set_style_font(doc.styles["Normal"], size=11)

    # Title (Arial Bold 14pt centered) — pandoc maps an H1 that's the
    # document's own <title>/first heading onto "Title" in some paths,
    # "Heading 1" in others; set both to the template spec so either maps
    # correctly.
    if "Title" in doc.styles:
        _set_style_font(doc.styles["Title"], size=14, bold=True)
        doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _set_style_font(doc.styles["Heading 1"], size=12, bold=True)
    _set_style_font(doc.styles["Heading 2"], size=12, bold=False)
    _set_style_font(doc.styles["Heading 3"], size=12, bold=False, italic=True)

    # Seed one paragraph per style so pandoc's reference-doc style
    # discovery (which only picks up styles actually used in the doc)
    # finds all of them.
    doc.add_paragraph("Title placeholder", style="Title")
    doc.add_paragraph("Heading 1 placeholder", style="Heading 1")
    doc.add_paragraph("Heading 2 placeholder", style="Heading 2")
    doc.add_paragraph("Heading 3 placeholder", style="Heading 3")
    doc.add_paragraph("Body text placeholder — Arial 11pt per template spec.")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
